from pathlib import Path
from datetime import datetime
import re
import unicodedata
import zipfile

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE


# ============================================================
# 1. CONFIGURAÇÕES — ALTERAR AQUI
# ============================================================

PASTA_BASE = Path(
    r"C:\Users\pedromaia\Downloads\EXTRAÇÃO_REAL_FINAL"
)

PASTA_ANALISAR = PASTA_BASE / "14_REL_AP_pdf"

EXCEL_SAIDA = PASTA_BASE / f"sugestoes_categorias_ficheiros_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

# True = analisa ficheiros dentro de subpastas.
# False = analisa apenas ficheiros diretamente dentro da pasta.
INCLUIR_SUBPASTAS = False

# O código ignora sempre estas extensões.
EXTENSOES_IGNORAR = {".txt"}

# None = analisa todas as páginas dos PDFs.
# Para acelerar, podes colocar 10, 20, 30, etc.
MAX_PAGINAS_PDF = None

# Limite de caracteres usados para análise textual.
MAX_CARACTERES_ANALISE = 300000

# Normalmente deixar False para o Excel não ficar pesado.
GUARDAR_AMOSTRA_TEXTO = False
TAMANHO_AMOSTRA_TEXTO = 1500


# ============================================================
# 2. CATEGORIAS
# ============================================================

CATEGORIAS = {
    "BBL": "Bibliografia",
    "SIG": "Dados SIG",
    "DES": "Desenho",
    "DAD": "Documentação Administrativa",
    "FTG": "Fotografia",
    "LIS": "Lista",
    "MRZ": "Matriz",
    "REL": "Relatório de projeto",
    "TBL": "Tabela",
    "VID": "Vídeo",
    "3D": "Modelo / levantamento 3D",
    "CMP": "Ficheiro compactado",
    "nd": "Não determinado automaticamente",
}


# ============================================================
# 3. CRITÉRIOS DE CLASSIFICAÇÃO
# ============================================================

REGRAS = {
    "REL": {
        "extensoes": {".pdf", ".docx", ".doc"},
        "nome": [
            "relatorio", "relatório", "rel", "final", "preliminar",
            "intervencao", "intervenção", "arqueologica", "arqueológica",
            "sondagem", "acompanhamento", "escavacao", "escavação"
        ],
        "texto_forte": [
            "relatório final", "relatorio final",
            "relatório preliminar", "relatorio preliminar",
            "relatório técnico", "relatorio tecnico",
            "relatório de trabalhos arqueológicos",
            "relatorio de trabalhos arqueologicos",
            "intervenção arqueológica", "intervencao arqueologica",
            "trabalhos arqueológicos", "trabalhos arqueologicos",
            "acompanhamento arqueológico", "acompanhamento arqueologico",
            "escavação arqueológica", "escavacao arqueologica",
            "sondagem arqueológica", "sondagem arqueologica"
        ],
        "texto_medio": [
            "ficha técnica", "ficha tecnica",
            "introdução", "introducao",
            "metodologia",
            "localização", "localizacao",
            "enquadramento",
            "resultados",
            "conclusão", "conclusao",
            "bibliografia"
        ],
        "texto_fraco": [
            "objetivos", "objectivos",
            "descrição dos trabalhos", "descricao dos trabalhos",
            "equipa técnica", "equipa tecnica"
        ],
    },

    "BBL": {
        "extensoes": {".pdf", ".docx", ".doc"},
        "nome": [
            "bibliografia", "artigo", "publicacao", "publicação",
            "paper", "revista", "livro", "capitulo", "capítulo",
            "separata", "monografia"
        ],
        "texto_forte": [
            "artigo científico", "artigo cientifico",
            "revista científica", "revista cientifica",
            "isbn", "issn", "doi",
            "bibliografia anotada",
            "referências bibliográficas", "referencias bibliograficas"
        ],
        "texto_medio": [
            "bibliografia",
            "referências", "referencias",
            "bibliographic references",
            "editora", "volume", "pp."
        ],
        "texto_fraco": [
            "autor", "autores", "editor", "editores",
            "páginas", "paginas"
        ],
    },

    "TBL": {
        "extensoes": {".xlsx", ".xls", ".csv", ".ods"},
        "nome": [
            "tabela", "base de dados", "bd", "database",
            "xls", "xlsx", "csv", "registo", "ficha",
            "campos", "dados"
        ],
        "texto_forte": [
            "tabela", "base de dados",
            "folha de cálculo", "folha de calculo",
            "registo em tabela",
            "campos da tabela",
            "estrutura da base de dados",
            "ficha de registo"
        ],
        "texto_medio": [
            "linha", "coluna", "campo", "registo",
            "código", "codigo", "id"
        ],
        "texto_fraco": [
            "n.º", "nº", "designação", "designacao"
        ],
    },

    "LIS": {
        "extensoes": {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".ods"},
        "nome": [
            "lista", "listagem", "inventario", "inventário",
            "espolio", "espólio", "materiais",
            "catalogo", "catálogo"
        ],
        "texto_forte": [
            "lista de materiais",
            "listagem de materiais",
            "inventário de espólio", "inventario de espolio",
            "catálogo de materiais", "catalogo de materiais",
            "espólio arqueológico", "espolio arqueologico",
            "número de inventário", "numero de inventario"
        ],
        "texto_medio": [
            "lista", "listagem", "inventário", "inventario",
            "espólio", "espolio", "materiais arqueológicos",
            "materiais arqueologicos"
        ],
        "texto_fraco": [
            "fragmento", "cerâmica", "ceramica",
            "vidro", "metal", "osso"
        ],
    },

    "FTG": {
        "extensoes": {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"},
        "nome": [
            "fotografia", "fotografico", "fotográfico",
            "foto", "fotos", "imagem", "imagens",
            "registo fotografico", "registo fotográfico",
            "album", "álbum"
        ],
        "texto_forte": [
            "registo fotográfico", "registo fotografico",
            "levantamento fotográfico", "levantamento fotografico",
            "documentação fotográfica", "documentacao fotografica",
            "álbum fotográfico", "album fotografico"
        ],
        "texto_medio": [
            "fotografia", "fotografias", "imagem", "imagens"
        ],
        "texto_fraco": [
            "figura", "registo de imagem"
        ],
    },

    "DES": {
        "extensoes": {".dwg", ".dxf", ".dgn", ".svg"},
        "nome": [
            "desenho", "planta", "plantas", "cad", "dwg", "dxf",
            "alcado", "alçado", "corte", "perfil",
            "levantamento grafico", "levantamento gráfico"
        ],
        "texto_forte": [
            "desenho arqueológico", "desenho arqueologico",
            "desenho técnico", "desenho tecnico",
            "levantamento gráfico", "levantamento grafico",
            "autocad", "ficheiro dwg", "ficheiro dxf"
        ],
        "texto_medio": [
            "planta", "plantas",
            "corte estratigráfico", "corte estratigrafico",
            "perfil estratigráfico", "perfil estratigrafico",
            "alçado", "alcado"
        ],
        "texto_fraco": [
            "escala gráfica", "escala grafica",
            "implantação", "implantacao",
            "desenho", "corte", "perfil"
        ],
    },

    "SIG": {
        "extensoes": {".shp", ".kml", ".kmz", ".gpkg", ".geojson", ".gdb", ".qgz", ".qgs"},
        "nome": [
            "sig", "gis", "qgis", "arcgis",
            "shp", "shape", "shapefile",
            "kmz", "kml", "gpkg", "geopackage",
            "geodatabase", "geojson"
        ],
        "texto_forte": [
            "sistema de informação geográfica",
            "sistema de informacao geografica",
            "dados sig", "qgis", "arcgis",
            "shapefile", "geopackage",
            "geodatabase", "base cartográfica", "base cartografica"
        ],
        "texto_medio": [
            "coordenadas", "georreferenciação", "georreferenciacao",
            "etrs89", "pt-tm06", "pttm06", "wgs84", "datum"
        ],
        "texto_fraco": [
            "cartografia",
            "coordenadas geográficas", "coordenadas geograficas",
            "georreferenciado", "georreferenciada"
        ],
    },

    "MRZ": {
        "extensoes": {".tif", ".tiff", ".ecw", ".asc", ".vrt"},
        "nome": [
            "matriz", "raster", "ortofoto", "ortofotomapa",
            "mdt", "mde", "dem", "dtm",
            "cartografia", "georreferenciada", "georreferenciado"
        ],
        "texto_forte": [
            "matriz raster",
            "modelo digital do terreno",
            "modelo digital de terreno",
            "modelo digital de elevação",
            "modelo digital de elevacao",
            "cartografia georreferenciada",
            "imagem georreferenciada",
            "ortofoto", "ortofotomapa"
        ],
        "texto_medio": [
            "raster", "georreferenciada",
            "georreferenciado", "cartografia"
        ],
        "texto_fraco": [
            "matriz", "modelo digital"
        ],
    },

    "DAD": {
        "extensoes": {".pdf", ".docx", ".doc", ".xlsx", ".xls"},
        "nome": [
            "administrativo", "administrativa",
            "contrato", "proposta", "orcamento", "orçamento",
            "fatura", "factura", "oficio", "ofício",
            "licenca", "licença", "despacho",
            "autorizacao", "autorização",
            "requerimento", "adjudicacao", "adjudicação"
        ],
        "texto_forte": [
            "contrato", "proposta de trabalhos", "proposta financeira",
            "orçamento", "orcamento", "fatura", "factura",
            "ofício", "oficio", "licença", "licenca",
            "despacho", "autorização", "autorizacao",
            "requerimento", "procedimento administrativo"
        ],
        "texto_medio": [
            "cliente", "adjudicação", "adjudicacao",
            "entidade adjudicante",
            "valor da proposta",
            "condições de pagamento", "condicoes de pagamento"
        ],
        "texto_fraco": [
            "nif", "morada fiscal", "assunto", "exmo", "exma"
        ],
    },

    "VID": {
        "extensoes": {".mp4", ".mov", ".avi", ".mkv", ".wmv", ".mpeg", ".mpg"},
        "nome": [
            "video", "vídeo", "filmagem", "mov", "mp4"
        ],
        "texto_forte": [],
        "texto_medio": [],
        "texto_fraco": [],
    },

    "3D": {
        "extensoes": {".e57", ".las", ".laz", ".obj", ".fbx", ".ply", ".rcp", ".rcs", ".psx"},
        "nome": [
            "3d", "e57", "laser", "scan", "scanner",
            "nuvem de pontos", "nuvem_pontos",
            "fotogrametria", "modelo 3d", "levantamento 3d"
        ],
        "texto_forte": [
            "nuvem de pontos", "levantamento laser",
            "laser scanning", "fotogrametria",
            "modelo tridimensional", "modelo 3d"
        ],
        "texto_medio": [
            "e57", "point cloud", "recap", "metashape"
        ],
        "texto_fraco": [],
    },

    "CMP": {
        "extensoes": {".zip", ".rar", ".7z"},
        "nome": [
            "zip", "rar", "compactado", "arquivo"
        ],
        "texto_forte": [],
        "texto_medio": [],
        "texto_fraco": [],
    },
}


# ============================================================
# 4. FUNÇÕES AUXILIARES
# ============================================================

def remover_acentos(texto: str) -> str:
    texto = unicodedata.normalize("NFD", str(texto))
    return "".join(c for c in texto if unicodedata.category(c) != "Mn")


def normalizar_texto(texto: str) -> str:
    texto = remover_acentos(str(texto))
    texto = texto.lower()
    texto = texto.replace("\xa0", " ")
    texto = texto.replace("\u200b", "")
    texto = texto.replace("\ufeff", "")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def limpar_para_excel(valor):
    if isinstance(valor, str):
        valor = ILLEGAL_CHARACTERS_RE.sub("", valor)

        if len(valor) > 32000:
            valor = valor[:32000] + "..."

    return valor


def limpar_dataframe_para_excel(df: pd.DataFrame) -> pd.DataFrame:
    return df.apply(lambda coluna: coluna.map(limpar_para_excel))


def listar_ficheiros(pasta: Path) -> list[Path]:
    if not pasta.exists():
        raise FileNotFoundError(f"Pasta não encontrada: {pasta}")

    if INCLUIR_SUBPASTAS:
        ficheiros = [p for p in pasta.rglob("*") if p.is_file()]
    else:
        ficheiros = [p for p in pasta.glob("*") if p.is_file()]

    return sorted([
        p for p in ficheiros
        if p.suffix.lower() not in EXTENSOES_IGNORAR
    ])


def contar_ocorrencias(texto_norm: str, termo: str) -> int:
    termo_norm = normalizar_texto(termo)

    if not termo_norm:
        return 0

    return texto_norm.count(termo_norm)


# ============================================================
# 5. EXTRAÇÃO DE CONTEÚDO
# ============================================================

def analise_base(caminho: Path) -> dict:
    return {
        "estado_leitura": "",
        "texto": "",
        "paginas_total": 0,
        "paginas_analisadas": 0,
        "caracteres_texto": 0,
        "paginas_com_texto": 0,
        "paginas_sem_texto": 0,
        "imagens_total": 0,
        "paginas_com_imagem": 0,
        "paginas_quase_so_imagem": 0,
        "media_cobertura_imagem": 0.0,
        "max_cobertura_imagem": 0.0,
        "ficheiro_apenas_imagem": False,
        "linhas_tabela_provavel": 0,
        "linhas_bibliografia_provavel": 0,
        "tem_ficha_tecnica": False,
        "tem_estrutura_relatorio": False,
        "folhas_excel": 0,
        "linhas_excel": 0,
        "colunas_excel": 0,
    }


def analisar_pdf(caminho: Path) -> dict:
    resultado = analise_base(caminho)
    resultado["estado_leitura"] = "pdf_lido"

    try:
        doc = fitz.open(str(caminho))
        total_paginas = len(doc)

        resultado["paginas_total"] = total_paginas

        if total_paginas == 0:
            resultado["estado_leitura"] = "pdf_sem_paginas"
            return resultado

        limite = total_paginas if MAX_PAGINAS_PDF is None else min(MAX_PAGINAS_PDF, total_paginas)
        resultado["paginas_analisadas"] = limite

        textos = []
        coberturas = []

        for i in range(limite):
            pagina = doc[i]
            texto_pagina = pagina.get_text("text") or ""
            texto_norm = normalizar_texto(texto_pagina)

            textos.append(texto_pagina)

            if len(texto_norm) >= 80:
                resultado["paginas_com_texto"] += 1
            else:
                resultado["paginas_sem_texto"] += 1

            imagens = pagina.get_images(full=True)
            resultado["imagens_total"] += len(imagens)

            blocos = pagina.get_text("dict").get("blocks", [])
            pagina_area = max(1.0, float(pagina.rect.width * pagina.rect.height))

            area_imagens = 0.0
            maior_imagem = 0.0

            for bloco in blocos:
                if bloco.get("type") == 1:
                    bbox = bloco.get("bbox", None)
                    if bbox:
                        x0, y0, x1, y1 = bbox
                        area = max(0.0, float((x1 - x0) * (y1 - y0)))
                        area_imagens += area
                        maior_imagem = max(maior_imagem, area)

            cobertura = min(1.0, area_imagens / pagina_area)
            cobertura_maior = min(1.0, maior_imagem / pagina_area)

            coberturas.append(cobertura)

            if len(imagens) > 0 or cobertura >= 0.35:
                resultado["paginas_com_imagem"] += 1

            if cobertura_maior >= 0.70 and len(texto_norm) < 80:
                resultado["paginas_quase_so_imagem"] += 1

        doc.close()

        texto_total = "\n".join(textos)
        texto_total = texto_total[:MAX_CARACTERES_ANALISE]

        resultado["texto"] = texto_total
        resultado["caracteres_texto"] = len(normalizar_texto(texto_total))

        if coberturas:
            resultado["media_cobertura_imagem"] = round(sum(coberturas) / len(coberturas), 3)
            resultado["max_cobertura_imagem"] = round(max(coberturas), 3)

        if limite > 0:
            taxa_imagem = resultado["paginas_quase_so_imagem"] / limite
            if taxa_imagem >= 0.75 and resultado["caracteres_texto"] < 800:
                resultado["ficheiro_apenas_imagem"] = True

        preencher_metricas_texto(resultado)

        return resultado

    except Exception as erro:
        resultado["estado_leitura"] = f"erro_pdf: {erro}"
        return resultado


def analisar_docx(caminho: Path) -> dict:
    resultado = analise_base(caminho)
    resultado["estado_leitura"] = "docx_lido"

    try:
        doc = Document(str(caminho))

        partes = []

        for p in doc.paragraphs:
            partes.append(p.text)

        linhas_tabelas = 0

        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [cel.text for cel in linha.cells]
                partes.append(" | ".join(celulas))
                linhas_tabelas += 1

        texto_total = "\n".join(partes)
        texto_total = texto_total[:MAX_CARACTERES_ANALISE]

        resultado["texto"] = texto_total
        resultado["caracteres_texto"] = len(normalizar_texto(texto_total))
        resultado["linhas_tabela_provavel"] = linhas_tabelas

        preencher_metricas_texto(resultado)

        return resultado

    except Exception as erro:
        resultado["estado_leitura"] = f"erro_docx: {erro}"
        return resultado


def analisar_excel(caminho: Path) -> dict:
    resultado = analise_base(caminho)
    resultado["estado_leitura"] = "excel_lido"

    try:
        xls = pd.ExcelFile(caminho)
        resultado["folhas_excel"] = len(xls.sheet_names)

        textos = []
        total_linhas = 0
        max_colunas = 0

        for folha in xls.sheet_names[:10]:
            df = pd.read_excel(caminho, sheet_name=folha, nrows=200)
            total_linhas += len(df)
            max_colunas = max(max_colunas, len(df.columns))

            textos.append(f"FOLHA: {folha}")
            textos.append("COLUNAS: " + " | ".join(map(str, df.columns.tolist())))

            amostra = df.head(30).astype(str).to_string(index=False)
            textos.append(amostra)

        texto_total = "\n".join(textos)
        texto_total = texto_total[:MAX_CARACTERES_ANALISE]

        resultado["texto"] = texto_total
        resultado["caracteres_texto"] = len(normalizar_texto(texto_total))
        resultado["linhas_excel"] = total_linhas
        resultado["colunas_excel"] = max_colunas
        resultado["linhas_tabela_provavel"] = total_linhas

        preencher_metricas_texto(resultado)

        return resultado

    except Exception as erro:
        resultado["estado_leitura"] = f"erro_excel: {erro}"
        return resultado


def analisar_csv(caminho: Path) -> dict:
    resultado = analise_base(caminho)
    resultado["estado_leitura"] = "csv_lido"

    try:
        df = pd.read_csv(caminho, sep=None, engine="python", nrows=500)

        texto_total = []
        texto_total.append("COLUNAS: " + " | ".join(map(str, df.columns.tolist())))
        texto_total.append(df.head(50).astype(str).to_string(index=False))

        texto_total = "\n".join(texto_total)
        texto_total = texto_total[:MAX_CARACTERES_ANALISE]

        resultado["texto"] = texto_total
        resultado["caracteres_texto"] = len(normalizar_texto(texto_total))
        resultado["linhas_excel"] = len(df)
        resultado["colunas_excel"] = len(df.columns)
        resultado["linhas_tabela_provavel"] = len(df)

        preencher_metricas_texto(resultado)

        return resultado

    except Exception as erro:
        resultado["estado_leitura"] = f"erro_csv: {erro}"
        return resultado


def analisar_generico(caminho: Path) -> dict:
    resultado = analise_base(caminho)

    ext = caminho.suffix.lower()

    if ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}:
        resultado["estado_leitura"] = "imagem_identificada_por_extensao"
        resultado["ficheiro_apenas_imagem"] = True
        resultado["imagens_total"] = 1
    else:
        resultado["estado_leitura"] = "ficheiro_analisado_por_nome_e_extensao"

    return resultado


def analisar_ficheiro(caminho: Path) -> dict:
    ext = caminho.suffix.lower()

    if ext == ".pdf":
        return analisar_pdf(caminho)

    if ext == ".docx":
        return analisar_docx(caminho)

    if ext in {".xlsx", ".xls", ".ods"}:
        return analisar_excel(caminho)

    if ext == ".csv":
        return analisar_csv(caminho)

    return analisar_generico(caminho)


# ============================================================
# 6. MÉTRICAS TEXTUAIS
# ============================================================

def preencher_metricas_texto(resultado: dict) -> None:
    texto = resultado.get("texto", "")

    resultado["linhas_tabela_provavel"] += contar_linhas_tabela(texto)
    resultado["linhas_bibliografia_provavel"] = contar_linhas_bibliografia(texto)
    resultado["tem_ficha_tecnica"] = "ficha tecnica" in normalizar_texto(texto)
    resultado["tem_estrutura_relatorio"] = detectar_estrutura_relatorio(texto)


def contar_linhas_tabela(texto: str) -> int:
    total = 0

    for linha in str(texto).splitlines():
        l = linha.strip()

        if len(l) < 5:
            continue

        tem_separadores = len(re.findall(r"\s{2,}|;|\|", l)) >= 2
        tem_muitos_numeros = len(re.findall(r"\b\d+[.,]?\d*\b", l)) >= 3
        tem_tabs = "\t" in l

        if tem_separadores or tem_muitos_numeros or tem_tabs:
            total += 1

    return total


def contar_linhas_bibliografia(texto: str) -> int:
    total = 0

    for linha in str(texto).splitlines():
        l = linha.strip()

        if len(l) < 20:
            continue

        tem_ano = re.search(r"\b(18|19|20)\d{2}\b", l)
        tem_autor = re.search(
            r"[A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-Za-zÁÉÍÓÚÂÊÔÃÕÇáéíóúâêôãõç\-]+,\s+[A-Z]",
            l
        )
        tem_publicacao = re.search(
            r"\b(pp\.|vol\.|doi|isbn|issn|revista|editora|in:)\b",
            l,
            flags=re.IGNORECASE
        )

        if tem_ano and (tem_autor or tem_publicacao):
            total += 1

    return total


def detectar_estrutura_relatorio(texto: str) -> bool:
    texto_norm = normalizar_texto(texto)

    secoes = [
        "ficha tecnica",
        "introducao",
        "metodologia",
        "objetivos",
        "objectivos",
        "localizacao",
        "enquadramento",
        "trabalhos arqueologicos",
        "resultados",
        "conclusao",
        "bibliografia",
    ]

    encontradas = sum(1 for s in secoes if s in texto_norm)

    return encontradas >= 4


# ============================================================
# 7. CLASSIFICAÇÃO
# ============================================================

def adicionar_score(scores: dict, codigo: str, pontos: int, evidencia: str) -> None:
    scores[codigo]["score"] += pontos

    if evidencia and len(scores[codigo]["evidencias"]) < 15:
        scores[codigo]["evidencias"].append(evidencia)


def classificar_ficheiro(caminho: Path, analise: dict) -> dict:
    nome_norm = normalizar_texto(caminho.name)
    texto_norm = normalizar_texto(analise.get("texto", ""))
    ext = caminho.suffix.lower()

    scores = {
        codigo: {"score": 0, "evidencias": []}
        for codigo in REGRAS
    }

    # A) Pontuação por extensão
    for codigo, regra in REGRAS.items():
        if ext in regra["extensoes"]:
            adicionar_score(scores, codigo, 35, f"extensão {ext} compatível com {codigo}")

    # B) Pontuação pelo nome
    for codigo, regra in REGRAS.items():
        for termo in regra["nome"]:
            if normalizar_texto(termo) in nome_norm:
                adicionar_score(scores, codigo, 25, f"nome contém '{termo}'")

    # C) Pontuação pelo texto
    for codigo, regra in REGRAS.items():
        for termo in regra["texto_forte"]:
            n = contar_ocorrencias(texto_norm, termo)
            if n:
                adicionar_score(scores, codigo, 12 * min(n, 5), f"texto contém '{termo}' ({n}x)")

        for termo in regra["texto_medio"]:
            n = contar_ocorrencias(texto_norm, termo)
            if n:
                adicionar_score(scores, codigo, 6 * min(n, 5), f"texto contém '{termo}' ({n}x)")

        for termo in regra["texto_fraco"]:
            n = contar_ocorrencias(texto_norm, termo)
            if n:
                adicionar_score(scores, codigo, 2 * min(n, 5), f"texto contém '{termo}' ({n}x)")

    # D) Pontuação estrutural
    if analise["ficheiro_apenas_imagem"]:
        adicionar_score(scores, "FTG", 80, "ficheiro é imagem ou PDF praticamente só com imagem")

    if analise["tem_ficha_tecnica"]:
        adicionar_score(scores, "REL", 30, "tem ficha técnica")

    if analise["tem_estrutura_relatorio"]:
        adicionar_score(scores, "REL", 55, "tem estrutura típica de relatório")

    if analise["linhas_tabela_provavel"] >= 30:
        adicionar_score(scores, "TBL", 55, f"muitas linhas com estrutura tabular ({analise['linhas_tabela_provavel']})")
    elif analise["linhas_tabela_provavel"] >= 10:
        adicionar_score(scores, "TBL", 30, f"algumas linhas com estrutura tabular ({analise['linhas_tabela_provavel']})")

    if analise["linhas_bibliografia_provavel"] >= 8 and not analise["tem_ficha_tecnica"]:
        adicionar_score(scores, "BBL", 60, f"muitas referências bibliográficas prováveis ({analise['linhas_bibliografia_provavel']}) e sem ficha técnica")
    elif analise["linhas_bibliografia_provavel"] >= 4 and not analise["tem_ficha_tecnica"]:
        adicionar_score(scores, "BBL", 35, f"algumas referências bibliográficas prováveis ({analise['linhas_bibliografia_provavel']}) e sem ficha técnica")

    if analise["folhas_excel"] > 0 or analise["colunas_excel"] >= 3:
        adicionar_score(scores, "TBL", 45, "estrutura de folha de cálculo/tabela")

    # E) Decisão final
    ranking = sorted(scores.items(), key=lambda item: item[1]["score"], reverse=True)

    melhor_codigo = ranking[0][0]
    melhor_score = ranking[0][1]["score"]
    segundo_codigo = ranking[1][0]
    segundo_score = ranking[1][1]["score"]

    diferenca = melhor_score - segundo_score

    if melhor_score < 20:
        codigo_final = "nd"
        categoria_final = CATEGORIAS["nd"]
        confianca = "baixa"
        estado = "não determinado automaticamente"
        evidencias = ""
    else:
        codigo_final = melhor_codigo
        categoria_final = CATEGORIAS.get(melhor_codigo, melhor_codigo)

        if diferenca <= 8 and segundo_score >= 20:
            confianca = "baixa-média"
            estado = f"VALIDAR: categorias próximas ({melhor_codigo} e {segundo_codigo})"
        elif melhor_score >= 75 and diferenca >= 20:
            confianca = "alta"
            estado = "categoria sugerida automaticamente"
        elif melhor_score >= 40:
            confianca = "média"
            estado = "categoria sugerida automaticamente"
        else:
            confianca = "baixa-média"
            estado = "categoria sugerida automaticamente; validar"

        evidencias = "; ".join(ranking[0][1]["evidencias"])

    return {
        "codigo_categoria_sugerida": codigo_final,
        "categoria_sugerida": categoria_final,
        "confianca": confianca,
        "score_categoria": melhor_score,
        "estado_categoria": estado,
        "categoria_secundaria": segundo_codigo,
        "score_secundario": segundo_score,
        "evidencias": evidencias,
        "ranking_scores": "; ".join(
            f"{codigo}={info['score']}"
            for codigo, info in ranking
            if info["score"] > 0
        ),
    }


# ============================================================
# 8. FOLHA DE CRITÉRIOS
# ============================================================

def criar_folha_criterios() -> pd.DataFrame:
    dados = [
        {"Codigo": "REL", "Categoria": "Relatório de projeto", "Critérios": "Relatório final/preliminar/técnico, intervenção arqueológica, ficha técnica, introdução, metodologia, localização, resultados, conclusão."},
        {"Codigo": "BBL", "Categoria": "Bibliografia", "Critérios": "Artigos, publicações, ISBN, ISSN, DOI, referências bibliográficas, padrão autor/ano, normalmente sem ficha técnica."},
        {"Codigo": "TBL", "Categoria": "Tabela", "Critérios": "Extensões Excel/CSV, linhas e colunas, base de dados, tabela, registos, campos, estrutura tabular."},
        {"Codigo": "LIS", "Categoria": "Lista", "Critérios": "Lista, listagem, inventário, espólio, catálogo de materiais, número de inventário."},
        {"Codigo": "FTG", "Categoria": "Fotografia", "Critérios": "Imagem isolada, PDF quase só com imagem, fotografia, registo fotográfico, álbum fotográfico."},
        {"Codigo": "DES", "Categoria": "Desenho", "Critérios": "Plantas, CAD, DWG, DXF, cortes, perfis, alçados, desenho técnico ou arqueológico."},
        {"Codigo": "SIG", "Categoria": "Dados SIG", "Critérios": "SHP, KML, KMZ, GeoPackage, QGIS, ArcGIS, coordenadas, georreferenciação, ETRS89, PT-TM06, WGS84."},
        {"Codigo": "MRZ", "Categoria": "Matriz", "Critérios": "Raster, ortofoto, ortofotomapa, matriz raster, modelo digital do terreno/elevação, cartografia georreferenciada."},
        {"Codigo": "DAD", "Categoria": "Documentação Administrativa", "Critérios": "Contrato, proposta, orçamento, fatura, ofício, licença, despacho, autorização, requerimento, adjudicação."},
        {"Codigo": "VID", "Categoria": "Vídeo", "Critérios": "Extensões de vídeo, como .mp4, .mov, .avi, .mkv."},
        {"Codigo": "3D", "Categoria": "Modelo / levantamento 3D", "Critérios": "E57, LAS, LAZ, OBJ, FBX, PLY, RCP, RCS, nuvem de pontos, laser scanning, fotogrametria."},
        {"Codigo": "CMP", "Categoria": "Ficheiro compactado", "Critérios": "ZIP, RAR, 7Z ou ficheiros compactados."},
        {"Codigo": "nd", "Categoria": "Não determinado automaticamente", "Critérios": "Usado quando não há evidência suficiente para sugerir uma categoria."},
    ]

    return pd.DataFrame(dados)


# ============================================================
# 9. PROCESSO PRINCIPAL
# ============================================================

def main():
    ficheiros = listar_ficheiros(PASTA_ANALISAR)

    print(f"Ficheiros encontrados para análise: {len(ficheiros)}")
    print(f"Extensões ignoradas: {', '.join(sorted(EXTENSOES_IGNORAR))}")

    resultados = []

    for i, ficheiro in enumerate(ficheiros, start=1):
        print(f"[{i}/{len(ficheiros)}] A analisar: {ficheiro.name}")

        analise = analisar_ficheiro(ficheiro)
        classificacao = classificar_ficheiro(ficheiro, analise)

        resultados.append({
            "nome_ficheiro": ficheiro.name,
            "caminho_ficheiro": str(ficheiro),
            "extensao": ficheiro.suffix.lower(),
            "estado_leitura": analise["estado_leitura"],
            "codigo_categoria_sugerida": classificacao["codigo_categoria_sugerida"],
            "categoria_sugerida": classificacao["categoria_sugerida"],
            "confianca": classificacao["confianca"],
            "score_categoria": classificacao["score_categoria"],
            "estado_categoria": classificacao["estado_categoria"],
            "categoria_secundaria": classificacao["categoria_secundaria"],
            "score_secundario": classificacao["score_secundario"],
            "evidencias": classificacao["evidencias"],
            "ranking_scores": classificacao["ranking_scores"],
            "paginas_total": analise["paginas_total"],
            "paginas_analisadas": analise["paginas_analisadas"],
            "caracteres_texto": analise["caracteres_texto"],
            "imagens_total": analise["imagens_total"],
            "paginas_com_imagem": analise["paginas_com_imagem"],
            "paginas_quase_so_imagem": analise["paginas_quase_so_imagem"],
            "ficheiro_apenas_imagem": analise["ficheiro_apenas_imagem"],
            "linhas_tabela_provavel": analise["linhas_tabela_provavel"],
            "linhas_bibliografia_provavel": analise["linhas_bibliografia_provavel"],
            "tem_ficha_tecnica": analise["tem_ficha_tecnica"],
            "tem_estrutura_relatorio": analise["tem_estrutura_relatorio"],
            "folhas_excel": analise["folhas_excel"],
            "linhas_excel": analise["linhas_excel"],
            "colunas_excel": analise["colunas_excel"],
            "amostra_texto": analise["texto"][:TAMANHO_AMOSTRA_TEXTO] if GUARDAR_AMOSTRA_TEXTO else "",
        })

    df_resultados = pd.DataFrame(resultados)

    resumo_categorias = (
        df_resultados
        .groupby(["codigo_categoria_sugerida", "categoria_sugerida", "confianca"])
        .size()
        .reset_index(name="total_ficheiros")
        .sort_values(["codigo_categoria_sugerida", "confianca"])
    )

    criterios = criar_folha_criterios()

    resumo = pd.DataFrame([
        {"Indicador": "Pasta analisada", "Valor": str(PASTA_ANALISAR)},
        {"Indicador": "Total de ficheiros analisados", "Valor": len(ficheiros)},
        {"Indicador": "Incluiu subpastas", "Valor": INCLUIR_SUBPASTAS},
        {"Indicador": "Extensões ignoradas", "Valor": ", ".join(sorted(EXTENSOES_IGNORAR))},
        {"Indicador": "Categorias possíveis", "Valor": ", ".join(CATEGORIAS.keys())},
        {"Indicador": "Ficheiros alterados", "Valor": "Nenhum ficheiro foi alterado"},
        {"Indicador": "Ação realizada", "Valor": "Apenas criação de Excel com sugestões de categoria"},
    ])

    df_resultados = limpar_dataframe_para_excel(df_resultados)
    resumo_categorias = limpar_dataframe_para_excel(resumo_categorias)
    criterios = limpar_dataframe_para_excel(criterios)
    resumo = limpar_dataframe_para_excel(resumo)

    with pd.ExcelWriter(EXCEL_SAIDA, engine="openpyxl") as writer:
        df_resultados.to_excel(writer, index=False, sheet_name="Categorias_sugeridas")
        resumo_categorias.to_excel(writer, index=False, sheet_name="Resumo_categorias")
        criterios.to_excel(writer, index=False, sheet_name="Criterios_usados")
        resumo.to_excel(writer, index=False, sheet_name="Resumo")

    print("\nProcesso concluído.")
    print(f"Excel criado em: {EXCEL_SAIDA}")
    print("Nenhum ficheiro foi alterado.")


if __name__ == "__main__":
    main()