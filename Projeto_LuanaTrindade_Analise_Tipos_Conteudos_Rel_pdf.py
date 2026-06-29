import os
import re
import zipfile
import unicodedata
from pathlib import Path
from datetime import datetime

import pandas as pd

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Alignment
except ImportError:
    load_workbook = None
    Font = None
    Alignment = None

try:
    from PIL import Image, ExifTags
except ImportError:
    Image = None
    ExifTags = None


# ==================================================
# CONFIGURAÇÃO
# ==================================================

# ALTERAR ESTA LINHA PARA A PASTA QUE QUERES ANALISAR
PASTA_BASE = r"\\rtvfserver\Dados\06_FOT_ORIGINAL"

# O Excel será criado na mesma pasta onde está este ficheiro Python
PASTA_DO_CODIGO = Path(__file__).parent
FICHEIRO_SAIDA = PASTA_DO_CODIGO / "Relatorio_Tipos_e_Conteudos_Luana.xlsx"

# Limites de leitura para evitar lentidão em pastas grandes
MAX_PAGINAS_PDF = 4
MAX_CARACTERES_TEXTO = 12000
MAX_FOLHAS_EXCEL = 5
MAX_LINHAS_EXCEL = 40
MAX_FICHEIROS_ZIP = 80


# ==================================================
# EXTENSÕES COM CLASSIFICAÇÃO PRIORITÁRIA
# ==================================================

VIDEO_EXTENSOES = {
    ".mov", ".mp4", ".avi", ".mkv", ".wmv", ".flv", ".webm", ".m4v"
}

AUDIO_EXTENSOES = {
    ".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a", ".wma"
}


# ==================================================
# TIPOS DE FICHEIRO POR EXTENSÃO
# ==================================================

TIPOS_POR_EXTENSAO = {
    ".pdf": "Documento PDF",
    ".doc": "Documento Word antigo",
    ".docx": "Documento Word",
    ".xls": "Folha de cálculo Excel antiga",
    ".xlsx": "Folha de cálculo Excel",
    ".xlsm": "Folha de cálculo Excel com macros",
    ".csv": "Tabela CSV",
    ".txt": "Texto simples",
    ".rtf": "Documento RTF",

    ".jpg": "Imagem JPEG",
    ".jpeg": "Imagem JPEG",
    ".png": "Imagem PNG",
    ".tif": "Imagem TIFF",
    ".tiff": "Imagem TIFF",
    ".bmp": "Imagem BMP",
    ".gif": "Imagem GIF",

    ".zip": "Arquivo ZIP",
    ".rar": "Arquivo RAR",
    ".7z": "Arquivo 7-Zip",

    ".kmz": "Ficheiro geográfico KMZ",
    ".kml": "Ficheiro geográfico KML",
    ".shp": "Shapefile GIS",
    ".dbf": "Tabela DBF / componente Shapefile",
    ".shx": "Índice Shapefile",
    ".prj": "Sistema de coordenadas Shapefile",
    ".geojson": "GeoJSON",
    ".json": "JSON",
    ".gpkg": "GeoPackage GIS",

    ".dwg": "Desenho CAD",
    ".dxf": "Desenho CAD DXF",

    ".obj": "Modelo 3D OBJ",
    ".ply": "Modelo 3D / nuvem de pontos PLY",
    ".las": "Nuvem de pontos LAS",
    ".laz": "Nuvem de pontos LAZ",
    ".e57": "Nuvem de pontos E57",

    ".mov": "Vídeo MOV",
    ".mp4": "Vídeo MP4",
    ".avi": "Vídeo AVI",
    ".mkv": "Vídeo MKV",
    ".wmv": "Vídeo WMV",
    ".flv": "Vídeo FLV",
    ".webm": "Vídeo WEBM",
    ".m4v": "Vídeo M4V",

    ".mp3": "Áudio MP3",
    ".wav": "Áudio WAV",
    ".aac": "Áudio AAC",
    ".flac": "Áudio FLAC",
    ".ogg": "Áudio OGG",
    ".m4a": "Áudio M4A",
    ".wma": "Áudio WMA"
}


# ==================================================
# PALAVRAS-CHAVE PARA INFERIR CONTEÚDO
# ==================================================

CATEGORIAS_CONTEUDO = {
    "Relatório de projeto / intervenção arqueológica": [
        "relatório", "relatorio", "relatório final", "relatorio final",
        "projeto", "projecto", "intervenção arqueológica", "intervencao arqueologica",
        "acompanhamento arqueológico", "acompanhamento arqueologico",
        "escavação", "escavacao", "sondagem", "sondagens",
        "diagnóstico arqueológico", "diagnostico arqueologico",
        "prospeção", "prospecção", "prospeccao",
        "memória descritiva", "memoria descritiva",
        "plano de trabalhos", "trabalhos arqueológicos", "trabalhos arqueologicos",
        "património arqueológico", "patrimonio arqueologico"
    ],

    "Lista / inventário de espólio": [
        "espólio", "espolio", "inventário", "inventario",
        "lista de espólio", "lista de espolio",
        "materiais arqueológicos", "materiais arqueologicos",
        "cerâmica", "ceramica", "lítico", "litico",
        "fauna", "metal", "vidro", "osso",
        "fragmento", "fragmentos",
        "unidade estratigráfica", "unidade estratigrafica",
        "ue", "contexto", "saco", "caixa", "etiqueta",
        "número de inventário", "numero de inventario",
        "n.º inventário", "nº inventário"
    ],

    "Fotografia / registo fotográfico": [
        "fotografia", "fotografias", "foto", "fotos",
        "registo fotográfico", "registo fotografico",
        "imagem", "imagens",
        "dsc", "img_", "jpeg", "jpg",
        "canon", "nikon", "sony", "olympus", "fujifilm",
        "câmara", "camara", "exif"
    ],

    "Fotogrametria / modelo 3D": [
        "fotogrametria", "photogrammetry",
        "agisoft", "metashape", "photoscan",
        "nuvem de pontos", "point cloud", "dense cloud",
        "mesh", "malha", "modelo 3d", "modelo tridimensional",
        "textura", "texture",
        "ortomosaico", "orthomosaic",
        "ortofoto", "ortofotografia",
        "dem", "mdt", "mde", "mns",
        "obj", "ply", "las", "laz", "e57"
    ],

    "Matriz / raster / cartografia georreferenciada": [
        "matriz", "raster", "geotiff", "geo tiff",
        "tiff georreferenciado", "imagem georreferenciada",
        "dem", "mdt", "mde", "mns",
        "hillshade", "declive", "altimetria",
        "ortofoto", "ortomosaico",
        "grid", "carta", "mapa", "cartografia",
        "epsg", "coordenadas", "georreferenciação", "georreferenciacao"
    ],

    "Dados SIG / GIS": [
        "sig", "gis", "shapefile", "geojson", "geopackage",
        "kml", "kmz", "layer", "camada", "feature",
        "qgis", "arcgis", "arcmap", "arcpro",
        "latitude", "longitude", "coordenadas",
        "epsg", "datum", "projeção", "projecao"
    ],

    "Desenho / planta / CAD": [
        "planta", "desenho", "cad", "autocad",
        "dwg", "dxf", "corte", "alçado", "alcado",
        "perfil", "levantamento", "escala",
        "implantação", "implantacao", "topografia"
    ],

    "Documentação administrativa": [
        "contrato", "fatura", "factura", "recibo",
        "orçamento", "orcamento", "proposta",
        "adjudicação", "adjudicacao",
        "ofício", "oficio", "ata", "email",
        "cronograma", "pagamento", "cliente", "fornecedor"
    ],

    "Tabela / base de dados": [
        "tabela", "base de dados", "registo", "registos",
        "lista", "código", "codigo", "id",
        "nome", "data", "descrição", "descricao",
        "observações", "observacoes", "categoria",
        "quantidade", "valor", "estado"
    ],

    "Bibliografia / artigo / publicação": [
        "bibliografia", "referências", "referencias",
        "artigo", "publicação", "publicacao",
        "revista", "doi", "issn", "isbn",
        "capítulo", "capitulo", "autor", "autores"
    ]
}


# ==================================================
# FUNÇÕES AUXILIARES
# ==================================================

def normalizar_texto(texto):
    texto = str(texto).lower()
    texto = unicodedata.normalize("NFD", texto)
    texto = "".join(c for c in texto if unicodedata.category(c) != "Mn")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def palavra_existe_no_texto(palavra, texto):
    palavra_norm = normalizar_texto(palavra)

    if not palavra_norm:
        return False

    # Para expressões com várias palavras ou termos com símbolos, usa procura direta
    if " " in palavra_norm or "_" in palavra_norm or "." in palavra_norm or "/" in palavra_norm:
        return palavra_norm in texto

    # Para palavras simples, usa limites de palavra para evitar erros como:
    # "id" aparecer dentro de "video"
    padrao = r"\b" + re.escape(palavra_norm) + r"\b"
    return re.search(padrao, texto) is not None


def encurtar(texto, limite=700):
    texto = re.sub(r"\s+", " ", str(texto)).strip()

    if len(texto) <= limite:
        return texto

    return texto[:limite].rstrip() + "..."


def formatar_tamanho(bytes_size):
    if bytes_size is None or pd.isna(bytes_size):
        return "Sem dados"

    bytes_size = int(bytes_size)

    if bytes_size < 1024:
        return f"{bytes_size} B"
    elif bytes_size < 1024 ** 2:
        return f"{bytes_size / 1024:.2f} KB"
    elif bytes_size < 1024 ** 3:
        return f"{bytes_size / (1024 ** 2):.2f} MB"
    else:
        return f"{bytes_size / (1024 ** 3):.2f} GB"


def identificar_tipo(extensao):
    extensao = extensao.lower()

    if extensao == "":
        return "Sem extensão"

    return TIPOS_POR_EXTENSAO.get(extensao, "Tipo não identificado")


def detectar_assinatura_interna(caminho):
    try:
        with open(caminho, "rb") as f:
            cabecalho = f.read(12)

        if cabecalho.startswith(b"%PDF"):
            return "Assinatura PDF"
        if cabecalho.startswith(b"\xFF\xD8\xFF"):
            return "Assinatura JPEG"
        if cabecalho.startswith(b"\x89PNG"):
            return "Assinatura PNG"
        if cabecalho.startswith(b"II*\x00") or cabecalho.startswith(b"MM\x00*"):
            return "Assinatura TIFF"
        if cabecalho.startswith(b"PK"):
            return "Assinatura ZIP/Office/KMZ"
        if cabecalho.startswith(b"{") or cabecalho.startswith(b"["):
            return "Possível JSON/texto estruturado"

        return "Assinatura não identificada"

    except Exception:
        return "Não foi possível ler assinatura"


def classificar_conteudo(texto_base):
    texto_norm = normalizar_texto(texto_base)
    resultados = []

    for categoria, palavras in CATEGORIAS_CONTEUDO.items():
        encontradas = []

        for palavra in palavras:
            if palavra_existe_no_texto(palavra, texto_norm):
                encontradas.append(palavra)

        if encontradas:
            resultados.append({
                "categoria": categoria,
                "score": len(encontradas),
                "evidencias": encontradas[:12]
            })

    if not resultados:
        return "Não determinado automaticamente", "Baixa", ""

    resultados = sorted(resultados, key=lambda x: x["score"], reverse=True)

    melhor = resultados[0]
    segunda = resultados[1] if len(resultados) > 1 else None

    if segunda and segunda["score"] >= max(2, melhor["score"] - 1):
        categoria_final = f"{melhor['categoria']} / {segunda['categoria']}"
        score_final = melhor["score"] + segunda["score"]
        evidencias = melhor["evidencias"] + segunda["evidencias"]
    else:
        categoria_final = melhor["categoria"]
        score_final = melhor["score"]
        evidencias = melhor["evidencias"]

    if score_final >= 7:
        confianca = "Alta"
    elif score_final >= 3:
        confianca = "Média"
    else:
        confianca = "Baixa"

    return categoria_final, confianca, ", ".join(evidencias[:15])


def criar_frase_resumo(tipo_tecnico, categoria, detalhe, evidencias):
    if categoria == "Não determinado automaticamente":
        frase = f"{tipo_tecnico}; o conteúdo não foi determinado automaticamente com segurança."
    else:
        frase = f"{tipo_tecnico}; o conteúdo aparenta corresponder a {categoria.lower()}."

    if detalhe:
        frase += f" {detalhe}"

    if evidencias:
        frase += f" Pistas usadas: {evidencias}."

    return encurtar(frase, 700)


# ==================================================
# EXTRAÇÃO DE INFORMAÇÃO INTERNA POR TIPO DE FICHEIRO
# ==================================================

def extrair_texto_pdf(caminho):
    if PdfReader is None:
        return "", "PDF não analisado internamente porque a biblioteca pypdf não está instalada.", "pypdf não instalado"

    try:
        reader = PdfReader(str(caminho))

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return "", "PDF encriptado ou protegido.", "PDF protegido"

        total_paginas = len(reader.pages)
        texto = ""

        for pagina in reader.pages[:MAX_PAGINAS_PDF]:
            texto += "\n" + (pagina.extract_text() or "")

        metadata = reader.metadata
        meta_texto = ""

        if metadata:
            meta_texto = " ".join(
                str(v) for v in metadata.values()
                if v is not None
            )

        texto_final = (meta_texto + "\n" + texto).strip()

        if texto_final:
            detalhe = f"Foram analisadas as primeiras {min(MAX_PAGINAS_PDF, total_paginas)} página(s) de um total de {total_paginas}."
        else:
            detalhe = f"O PDF tem {total_paginas} página(s), mas não foi encontrado texto extraível; pode ser uma digitalização ou imagem."

        return texto_final[:MAX_CARACTERES_TEXTO], detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar extrair texto do PDF.", str(erro)


def extrair_texto_docx(caminho):
    if Document is None:
        return "", "DOCX não analisado internamente porque a biblioteca python-docx não está instalada.", "python-docx não instalado"

    try:
        doc = Document(str(caminho))

        partes = []

        for p in doc.paragraphs:
            if p.text.strip():
                partes.append(p.text.strip())

        for tabela in doc.tables[:5]:
            for linha in tabela.rows[:20]:
                valores = [
                    celula.text.strip()
                    for celula in linha.cells
                    if celula.text.strip()
                ]

                if valores:
                    partes.append(" | ".join(valores))

        texto = "\n".join(partes)

        detalhe = f"O documento tem {len(doc.paragraphs)} parágrafo(s) e {len(doc.tables)} tabela(s)."

        return texto[:MAX_CARACTERES_TEXTO], detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar extrair texto do documento Word.", str(erro)


def extrair_texto_excel(caminho):
    if load_workbook is None:
        return "", "Excel não analisado internamente porque a biblioteca openpyxl não está instalada.", "openpyxl não instalado"

    try:
        wb = load_workbook(str(caminho), read_only=True, data_only=True)

        partes = []
        nomes_folhas = wb.sheetnames

        partes.append("Folhas: " + ", ".join(nomes_folhas))

        for nome_folha in nomes_folhas[:MAX_FOLHAS_EXCEL]:
            ws = wb[nome_folha]
            partes.append(f"Folha: {nome_folha}")

            linhas_lidas = 0

            for linha in ws.iter_rows(max_row=MAX_LINHAS_EXCEL, values_only=True):
                valores = [
                    str(valor).strip()
                    for valor in linha
                    if valor is not None and str(valor).strip()
                ]

                if valores:
                    partes.append(" | ".join(valores[:20]))
                    linhas_lidas += 1

                if linhas_lidas >= MAX_LINHAS_EXCEL:
                    break

        wb.close()

        texto = "\n".join(partes)

        detalhe = f"A folha de cálculo tem {len(nomes_folhas)} folha(s): {', '.join(nomes_folhas[:8])}."

        return texto[:MAX_CARACTERES_TEXTO], detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar ler conteúdo do Excel.", str(erro)


def extrair_texto_simples(caminho):
    encodings = ["utf-8-sig", "utf-8", "latin-1", "cp1252"]

    for enc in encodings:
        try:
            with open(caminho, "r", encoding=enc, errors="ignore") as f:
                texto = f.read(MAX_CARACTERES_TEXTO)

            detalhe = f"O ficheiro foi lido como texto com codificação {enc}."
            return texto, detalhe, ""

        except Exception:
            continue

    return "", "Não foi possível ler o ficheiro como texto.", "Erro de codificação/leitura"


def extrair_texto_kml(caminho):
    try:
        texto, detalhe, erro = extrair_texto_simples(caminho)

        texto_sem_tags = re.sub(r"<[^>]+>", " ", texto)
        texto_sem_tags = re.sub(r"\s+", " ", texto_sem_tags)

        return texto_sem_tags[:MAX_CARACTERES_TEXTO], "Ficheiro KML com dados geográficos em XML.", erro

    except Exception as erro:
        return "", "Erro ao tentar ler KML.", str(erro)


def extrair_texto_kmz(caminho):
    try:
        partes = []

        with zipfile.ZipFile(caminho, "r") as z:
            nomes = z.namelist()
            partes.append("Ficheiros dentro do KMZ: " + ", ".join(nomes[:MAX_FICHEIROS_ZIP]))

            for nome in nomes:
                if nome.lower().endswith(".kml"):
                    with z.open(nome) as f:
                        conteudo = f.read(MAX_CARACTERES_TEXTO).decode("utf-8", errors="ignore")
                        conteudo = re.sub(r"<[^>]+>", " ", conteudo)
                        partes.append(conteudo)
                    break

        texto = "\n".join(partes)

        detalhe = "O KMZ foi analisado através dos nomes internos e do primeiro KML encontrado."

        return texto[:MAX_CARACTERES_TEXTO], detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar ler KMZ.", str(erro)


def extrair_texto_zip(caminho):
    try:
        with zipfile.ZipFile(caminho, "r") as z:
            nomes = z.namelist()

        texto = "Ficheiros dentro do ZIP: " + ", ".join(nomes[:MAX_FICHEIROS_ZIP])
        detalhe = f"O arquivo ZIP tem {len(nomes)} item(ns) interno(s), analisados pelo nome."

        return texto[:MAX_CARACTERES_TEXTO], detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar ler ZIP.", str(erro)


def extrair_info_imagem(caminho):
    if Image is None:
        return "", "Imagem não analisada internamente porque a biblioteca Pillow não está instalada.", "Pillow não instalado"

    try:
        with Image.open(caminho) as img:
            largura, altura = img.size
            formato = img.format
            modo = img.mode
            n_frames = getattr(img, "n_frames", 1)

            partes = [
                f"Imagem {formato}",
                f"dimensões {largura}x{altura}",
                f"modo {modo}",
                f"{n_frames} frame(s)"
            ]

            detalhe = f"A imagem tem {largura}x{altura}px, formato {formato} e modo {modo}."

            try:
                exif = img.getexif()

                if exif and ExifTags is not None:
                    tags = {
                        ExifTags.TAGS.get(k, k): v
                        for k, v in exif.items()
                    }

                    make = tags.get("Make", "")
                    model = tags.get("Model", "")
                    data_original = tags.get("DateTimeOriginal", "")

                    if make or model:
                        partes.append(f"câmara {make} {model}")
                        detalhe += f" Metadados EXIF indicam possível fotografia captada com {make} {model}."

                    if data_original:
                        partes.append(f"data original {data_original}")

                    if "GPSInfo" in tags:
                        partes.append("metadados GPS")
                        detalhe += " Contém possíveis metadados GPS."

            except Exception:
                pass

            try:
                geotiff_tags = {33550, 33922, 34735, 34736, 34737}
                tags_tiff = set(getattr(img, "tag_v2", {}).keys())

                if geotiff_tags.intersection(tags_tiff):
                    partes.append("geotiff raster georreferenciado matriz cartografia sig")
                    detalhe += " Foram detetadas tags compatíveis com GeoTIFF/raster georreferenciado."

            except Exception:
                pass

            texto = " ".join(partes)

            return texto, detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar analisar imagem.", str(erro)


def analisar_componentes_shapefile(caminho):
    try:
        componentes = sorted([
            p.suffix.lower()
            for p in caminho.parent.glob(caminho.stem + ".*")
        ])

        texto = f"Shapefile SIG GIS camada geográfica componentes: {' '.join(componentes)}"
        detalhe = f"Foram encontrados componentes associados ao Shapefile: {', '.join(componentes)}."

        return texto, detalhe, ""

    except Exception as erro:
        return "", "Erro ao tentar analisar componentes do Shapefile.", str(erro)


def analisar_formato_limitado(caminho, extensao, tipo_tecnico):
    texto = f"{caminho.name} {caminho.parent} {extensao} {tipo_tecnico}"
    detalhe = "A análise interna deste formato é limitada; a inferência baseou-se sobretudo no nome, caminho e extensão."

    return texto, detalhe, ""


# ==================================================
# ANÁLISE GERAL DE CADA FICHEIRO
# ==================================================

def analisar_conteudo_ficheiro(caminho):
    extensao = caminho.suffix.lower()
    tipo_tecnico = identificar_tipo(extensao)
    assinatura = detectar_assinatura_interna(caminho)

    # ==================================================
    # REGRA PRIORITÁRIA PARA VÍDEO
    # ==================================================
    if extensao in VIDEO_EXTENSOES:
        return {
            "Tipo identificado": tipo_tecnico,
            "Assinatura interna": assinatura,
            "Categoria provável do conteúdo": "Vídeo / registo audiovisual",
            "Confiança da classificação": "Alta",
            "Evidências encontradas": f"Extensão {extensao}; tipo técnico {tipo_tecnico}",
            "Descrição breve do conteúdo": (
                f"{tipo_tecnico}; ficheiro classificado como vídeo/registo audiovisual "
                f"com base na extensão {extensao}."
            ),
            "Amostra de texto/metadados": "",
            "Erro/limitação na análise interna": (
                "O conteúdo interno do vídeo não foi analisado; a classificação foi feita pela extensão."
            )
        }

    # ==================================================
    # REGRA PRIORITÁRIA PARA ÁUDIO
    # ==================================================
    if extensao in AUDIO_EXTENSOES:
        return {
            "Tipo identificado": tipo_tecnico,
            "Assinatura interna": assinatura,
            "Categoria provável do conteúdo": "Áudio / registo sonoro",
            "Confiança da classificação": "Alta",
            "Evidências encontradas": f"Extensão {extensao}; tipo técnico {tipo_tecnico}",
            "Descrição breve do conteúdo": (
                f"{tipo_tecnico}; ficheiro classificado como áudio/registo sonoro "
                f"com base na extensão {extensao}."
            ),
            "Amostra de texto/metadados": "",
            "Erro/limitação na análise interna": (
                "O conteúdo interno do áudio não foi analisado; a classificação foi feita pela extensão."
            )
        }

    texto_extraido = ""
    detalhe = ""
    erro_conteudo = ""

    if extensao == ".pdf":
        texto_extraido, detalhe, erro_conteudo = extrair_texto_pdf(caminho)

    elif extensao == ".docx":
        texto_extraido, detalhe, erro_conteudo = extrair_texto_docx(caminho)

    elif extensao in [".xlsx", ".xlsm"]:
        texto_extraido, detalhe, erro_conteudo = extrair_texto_excel(caminho)

    elif extensao in [".csv", ".txt", ".rtf", ".json", ".geojson"]:
        texto_extraido, detalhe, erro_conteudo = extrair_texto_simples(caminho)

    elif extensao == ".kml":
        texto_extraido, detalhe, erro_conteudo = extrair_texto_kml(caminho)

    elif extensao == ".kmz":
        texto_extraido, detalhe, erro_conteudo = extrair_texto_kmz(caminho)

    elif extensao == ".zip":
        texto_extraido, detalhe, erro_conteudo = extrair_texto_zip(caminho)

    elif extensao in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".gif"]:
        texto_extraido, detalhe, erro_conteudo = extrair_info_imagem(caminho)

    elif extensao in [".shp", ".dbf", ".shx", ".prj"]:
        texto_extraido, detalhe, erro_conteudo = analisar_componentes_shapefile(caminho)

    else:
        texto_extraido, detalhe, erro_conteudo = analisar_formato_limitado(caminho, extensao, tipo_tecnico)

    contexto_para_classificar = " ".join([
        caminho.name,
        str(caminho.parent),
        extensao,
        tipo_tecnico,
        assinatura,
        detalhe,
        texto_extraido
    ])

    categoria, confianca, evidencias = classificar_conteudo(contexto_para_classificar)

    descricao = criar_frase_resumo(
        tipo_tecnico=tipo_tecnico,
        categoria=categoria,
        detalhe=detalhe,
        evidencias=evidencias
    )

    return {
        "Tipo identificado": tipo_tecnico,
        "Assinatura interna": assinatura,
        "Categoria provável do conteúdo": categoria,
        "Confiança da classificação": confianca,
        "Evidências encontradas": evidencias,
        "Descrição breve do conteúdo": descricao,
        "Amostra de texto/metadados": encurtar(texto_extraido, 1000),
        "Erro/limitação na análise interna": erro_conteudo
    }


# ==================================================
# PERCORRER A PASTA
# ==================================================

colunas = [
    "Nome do ficheiro",
    "Extensão",
    "Tipo identificado",
    "Categoria provável do conteúdo",
    "Confiança da classificação",
    "Descrição breve do conteúdo",
    "Evidências encontradas",
    "Assinatura interna",
    "Amostra de texto/metadados",
    "Erro/limitação na análise interna",
    "Pasta",
    "Caminho completo",
    "Tamanho",
    "Tamanho em bytes",
    "Data de criação",
    "Data de modificação"
]

dados = []

pasta = Path(PASTA_BASE)

if not pasta.exists():
    raise FileNotFoundError(f"A pasta não existe: {PASTA_BASE}")

for raiz, diretorios, ficheiros in os.walk(pasta):
    for ficheiro in ficheiros:
        caminho = Path(raiz) / ficheiro

        try:
            estatisticas = caminho.stat()
            analise = analisar_conteudo_ficheiro(caminho)

            dados.append({
                "Nome do ficheiro": caminho.name,
                "Extensão": caminho.suffix.lower() if caminho.suffix else "Sem extensão",
                "Tipo identificado": analise["Tipo identificado"],
                "Categoria provável do conteúdo": analise["Categoria provável do conteúdo"],
                "Confiança da classificação": analise["Confiança da classificação"],
                "Descrição breve do conteúdo": analise["Descrição breve do conteúdo"],
                "Evidências encontradas": analise["Evidências encontradas"],
                "Assinatura interna": analise["Assinatura interna"],
                "Amostra de texto/metadados": analise["Amostra de texto/metadados"],
                "Erro/limitação na análise interna": analise["Erro/limitação na análise interna"],
                "Pasta": str(caminho.parent),
                "Caminho completo": str(caminho),
                "Tamanho": formatar_tamanho(estatisticas.st_size),
                "Tamanho em bytes": estatisticas.st_size,
                "Data de criação": datetime.fromtimestamp(estatisticas.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
                "Data de modificação": datetime.fromtimestamp(estatisticas.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            })

        except Exception as erro:
            dados.append({
                "Nome do ficheiro": ficheiro,
                "Extensão": Path(ficheiro).suffix.lower() if Path(ficheiro).suffix else "Sem extensão",
                "Tipo identificado": "Erro ao analisar",
                "Categoria provável do conteúdo": "Erro",
                "Confiança da classificação": "Baixa",
                "Descrição breve do conteúdo": "Não foi possível analisar este ficheiro.",
                "Evidências encontradas": "",
                "Assinatura interna": "Erro",
                "Amostra de texto/metadados": "",
                "Erro/limitação na análise interna": str(erro),
                "Pasta": raiz,
                "Caminho completo": str(caminho),
                "Tamanho": "Erro",
                "Tamanho em bytes": None,
                "Data de criação": "Erro",
                "Data de modificação": "Erro"
            })


df_ficheiros = pd.DataFrame(dados, columns=colunas)


# ==================================================
# RESUMOS
# ==================================================

df_resumo_extensao = (
    df_ficheiros
    .groupby(["Extensão", "Tipo identificado"], dropna=False)
    .agg(
        Quantidade=("Nome do ficheiro", "count"),
        Tamanho_Total_Bytes=("Tamanho em bytes", "sum")
    )
    .reset_index()
)

df_resumo_extensao["Tamanho total"] = df_resumo_extensao["Tamanho_Total_Bytes"].apply(formatar_tamanho)


df_resumo_categoria = (
    df_ficheiros
    .groupby(["Categoria provável do conteúdo", "Confiança da classificação"], dropna=False)
    .agg(
        Quantidade=("Nome do ficheiro", "count"),
        Tamanho_Total_Bytes=("Tamanho em bytes", "sum")
    )
    .reset_index()
)

df_resumo_categoria["Tamanho total"] = df_resumo_categoria["Tamanho_Total_Bytes"].apply(formatar_tamanho)


df_resumo_pasta = (
    df_ficheiros
    .groupby("Pasta", dropna=False)
    .agg(
        Quantidade_Ficheiros=("Nome do ficheiro", "count"),
        Tamanho_Total_Bytes=("Tamanho em bytes", "sum")
    )
    .reset_index()
)

df_resumo_pasta["Tamanho total"] = df_resumo_pasta["Tamanho_Total_Bytes"].apply(formatar_tamanho)


# ==================================================
# EXPORTAÇÃO PARA EXCEL
# ==================================================

with pd.ExcelWriter(FICHEIRO_SAIDA, engine="openpyxl") as writer:
    df_ficheiros.to_excel(writer, sheet_name="Ficheiros", index=False)
    df_resumo_extensao.to_excel(writer, sheet_name="Resumo por extensão", index=False)
    df_resumo_categoria.to_excel(writer, sheet_name="Resumo por categoria", index=False)
    df_resumo_pasta.to_excel(writer, sheet_name="Resumo por pasta", index=False)


# ==================================================
# FORMATAÇÃO DO EXCEL
# ==================================================

try:
    wb = load_workbook(FICHEIRO_SAIDA)

    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions

        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        for coluna in ws.columns:
            letra = coluna[0].column_letter
            maior = 0

            for celula in coluna:
                if celula.value is not None:
                    maior = max(maior, len(str(celula.value)))

            ws.column_dimensions[letra].width = min(maior + 2, 60)

    wb.save(FICHEIRO_SAIDA)

except Exception:
    pass


print("Análise concluída com sucesso!")
print(f"Total de ficheiros analisados: {len(df_ficheiros)}")
print("Excel criado em:")
print(FICHEIRO_SAIDA)